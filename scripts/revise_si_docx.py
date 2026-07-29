from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


ROOT = Path("/Users/yusong/Downloads/new")
INPUT = ROOT / "司.docx"
OUTPUT = ROOT / "司_导师意见修订版.docx"


def set_run_font(run, size: int | None = None, bold: bool | None = None) -> None:
    run.font.name = "宋体"
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_paragraph_text(paragraph, text: str) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_font(run, 12)


def find_paragraph(doc: Document, needle: str, start: int = 0):
    for paragraph in doc.paragraphs[start:]:
        if needle in paragraph.text:
            return paragraph
    raise ValueError(f"paragraph not found after {start}: {needle}")


def insert_paragraph_after(paragraph, text: str, style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    inserted = Paragraph(new_p, paragraph._parent)
    if style:
        inserted.style = style
    run = inserted.add_run(text)
    set_run_font(run, 12)
    return inserted


def format_caption(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        set_run_font(run, 10)


def insert_caption_after(paragraph, text: str):
    caption = insert_paragraph_after(paragraph, text)
    format_caption(caption)
    return caption


def insert_caption_before(paragraph, text: str):
    caption = insert_paragraph_before(paragraph, text)
    format_caption(caption)
    return caption


def insert_table_after(paragraph, rows: list[list[str]]):
    doc = paragraph.part.document
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1
            run = p.add_run(value)
            set_run_font(run, 8, bold=(r_idx == 0))
    paragraph._p.addnext(table._tbl)
    return table


def insert_paragraph_before(paragraph, text: str, style: str | None = None):
    inserted = paragraph.insert_paragraph_before(text)
    if style:
        inserted.style = style
    for run in inserted.runs:
        set_run_font(run, 12)
    return inserted


def insert_table_before(paragraph, rows: list[list[str]]):
    doc = paragraph.part.document
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1
            run = p.add_run(value)
            set_run_font(run, 8, bold=(r_idx == 0))
    paragraph._p.addprevious(table._tbl)
    return table


def replace_first_containing(doc: Document, old: str, new: str) -> None:
    paragraph = find_paragraph(doc, old)
    paragraph.text = paragraph.text.replace(old, new)
    for run in paragraph.runs:
        set_run_font(run, 12)


def replace_all_paragraph_text(doc: Document, old: str, new: str) -> None:
    for paragraph in doc.paragraphs:
        if old in paragraph.text:
            paragraph.text = paragraph.text.replace(old, new)
            for run in paragraph.runs:
                set_run_font(run, 12)


def renumber_chapter7_tables_for_new_sample(doc: Document) -> None:
    for old, new in [
        ("表 7-8", "表 7-9"),
        ("表 7-7", "表 7-8"),
        ("表 7-6", "表 7-7"),
        ("表 7-5", "表 7-6"),
        ("表 7-4", "表 7-5"),
        ("表 7-3", "表 7-4"),
        ("表 7-2", "表 7-3"),
    ]:
        replace_all_paragraph_text(doc, old, new)


def delete_existing_inserted_blocks(doc: Document) -> None:
    # Keep script idempotent enough for repeated local runs.
    markers = [
        "表 3-2 将本文处理的垂类场景",
        "从功能模块角度看，系统并不只包含笔记录入",
        "表 7-3 给出三类代表性记录",
        "表 B-2 进一步列出主要数据表",
        "本文中的“智能体”不指开放式自主规划智能体",
    ]
    # Full deletion of arbitrary following tables is intentionally avoided.
    # This script is expected to run from the original input file.
    for marker in markers:
        for p in doc.paragraphs:
            if marker in p.text:
                raise RuntimeError("The target document already appears revised; start from original 司.docx.")


def main() -> None:
    doc = Document(INPUT)
    delete_existing_inserted_blocks(doc)

    replacements = {
        "第三章 面向科研实验记录的需求分析与问题建模": "第三章 面向科研实验记录的需求分析与问题定义",
        "3.6 问题建模": "3.6 关键问题定义与设计约束提炼",
        "关键问题建模": "关键问题定义",
        "关键问题进行了建模": "关键问题进行了定义",
        "需求分析和问题建模": "需求分析和问题定义",
        "1.5 研究贡献与创新": "1.5 研究贡献与创新边界",
        "第二章 相关理论与技术基础": "第二章 相关理论与技术",
        "2.2 大语言模型与增强检索技术": "2.2 大语言模型与检索增强生成技术",
    }
    for old, new in replacements.items():
        for p in doc.paragraphs:
            if old in p.text:
                p.text = p.text.replace(old, new)
                for run in p.runs:
                    set_run_font(run, 12)
    renumber_chapter7_tables_for_new_sample(doc)

    # 3.4: biomedical vertical scenario mapping.
    anchor_34 = find_paragraph(doc, "每条问答记录应保存问题、回答、模式标识", start=180)
    p = insert_paragraph_after(
        anchor_34,
        "由于本文系统面向生物医学课题组场景，AI 功能不能仅停留在“接入大模型”层面，而需要对应到具体实验对象、记录形态和验证证据。表 3-2 将本文处理的垂类场景、原始记录形态、AI 处理目标、系统模块和后续验证材料对应起来。该表用于回答“系统到底解决了什么实验记录问题”：PCR、Western Blot、细胞活力检测等记录首先以结构化字段、正文和附件形式进入系统，随后被转化为实体、关系、资料块和问答日志，最后通过图谱核验、RAG 对照和审计记录验证处理结果。",
    )
    insert_table_after(
        p,
        [
            ["生物医学场景", "原始记录形态", "AI 处理目标", "系统落点", "后续验证证据"],
            ["PCR 条件优化", "实验类型、样本、试剂、退火温度、扩增结果和附件说明", "识别样本、试剂、条件和结果，回答条件依据及结果来源", "实验笔记模板、uses_reagent、uses_sample、produces_result、RAG 问答", "Q01、Q02、Q06、Q07、Q11 对照日志和图谱关系"],
            ["Western Blot 蛋白表达验证", "抗体/试剂、仪器、处理组、条带或表达变化结论", "结构化试剂、仪器和结果关系，支持来源追溯", "知识图谱抽取、图谱可视化、图谱增强 RAG", "Q09、Q10、Q13 对照日志和 53 条关系核验"],
            ["CCK-8/细胞活力检测", "细胞处理条件、检测步骤、读数要求、统计结论", "检索资料事实，关联处理组、仪器和实验结论", "资料审核入库、向量检索、图谱关系注入", "Q03、Q04、Q08、Q12 对照日志"],
            ["项目阶段复盘", "多条已审核笔记、知识文档、图谱关系和问答记录", "汇总指定时间范围内实验进展，生成可核验草稿", "固定任务型实验总结、周报、阶段报告和图谱概览", "第 7.6 节四类生成记录、来源笔记和来源关系"],
            ["项目知识追溯", "资料块、图谱关系、模型输入输出和审计日志", "保存每次问答的证据链，区分资料不足、关系不足和生成失败", "ai_query_logs、ai_experiment_runs、audit_logs", "日志 446 至 485、语料哈希、盲评表和解盲键"],
        ],
    )
    insert_caption_after(p, "表 3-2 生物医学实验场景与系统处理证据对应关系")

    # 5.5: fixed-task agent design table.
    heading_55 = find_paragraph(doc, "5.5 面向实验笔记的固定任务型辅助生成设计", start=250)
    p = insert_paragraph_after(
        heading_55,
        "本文中的“智能体”不指开放式自主规划智能体，而是指由系统预设输入、工具、提示模板和输出格式的固定任务型辅助生成单元。所有任务单元共享项目权限校验、已审核数据过滤、DeepSeek 受证据约束生成和来源日志保存四个步骤，区别在于任务目标、上下文组织方式和输出模板不同。表 5-1 给出系统当前设计的固定任务型智能体。",
    )
    insert_paragraph_after(
        p,
        "上述任务单元之间不进行自动协作或递归调用，也不允许模型自行选择数据库表或工具。系统先由后端根据权限和任务参数准备证据，再将证据输入 DeepSeek，最后将生成结果和来源编号写入 agent_generation_runs。该设计牺牲了开放智能体的灵活性，但更符合实验记录场景对可控性、可追溯性和人工复核的要求。",
    )
    insert_table_after(
        p,
        [
            ["固定任务型智能体", "输入范围", "调用工具或服务", "生成调用", "输出结果", "追溯与约束"],
            ["实验总结智能体", "指定项目、实验类型、时间范围内的已审核笔记、资料和关系", "笔记查询、资料查询、图谱关系查询", "根据已筛选证据生成结构化实验总结", "实验目的、材料条件、过程、结果和待复核问题", "保存来源笔记、资料和关系 ID；不回写原笔记"],
            ["周报生成智能体", "指定周内已审核笔记、项目资料和代表性关系", "时间范围过滤、项目权限校验、证据组装", "根据周内证据生成周报草稿", "本周进展、主要结果、异常、下周计划", "证据不足时必须说明；仅作为草稿"],
            ["阶段报告智能体", "较长时间范围内的项目笔记、资料和问答记录", "笔记聚合、资料列表、图谱上下文", "根据阶段证据生成阶段性报告", "阶段目标、完成实验、风险和后续安排", "结论必须能回到来源记录核验"],
            ["图谱概览智能体", "项目实体统计、关系类型统计和代表性关系", "图谱统计、关系筛选、来源编号映射", "将结构化图谱关系转写为自然语言说明", "实验对象分布、关键关系和来源笔记概览", "不新增关系，不推断未知事实"],
        ],
    )
    insert_caption_after(p, "表 5-1 固定任务型智能体设计")

    # 6.2: explicit module table.
    anchor_62 = find_paragraph(doc, "项目工作台截图", start=300)
    p = insert_paragraph_after(
        anchor_62,
        "从功能模块角度看，系统并不只包含笔记录入和关键词检索，而是覆盖从原始记录进入、可信审核、结构化转化、智能问答到结果评价的完整链路。表 6-1 列出主要模块及其输入输出。",
    )
    insert_table_after(
        p,
        [
            ["模块", "主要输入", "核心处理", "主要输出"],
            ["用户与项目权限模块", "用户账号、项目成员、项目角色", "登录认证、系统级与项目级权限判定", "可访问项目、读写审核管理权限"],
            ["实验笔记模块", "实验类型、结构化字段、正文、附件", "草稿保存、版本记录、提交审核、状态流转", "已审核笔记、版本历史、审核记录"],
            ["资料与附件模块", "笔记附件、项目知识文档", "文件存储、哈希记录、资料审核、下载留痕", "可追溯附件、待入库或已入库资料"],
            ["文档解析与向量入库模块", "已审核知识文档", "文本抽取、分块、嵌入、pgvector 写入", "项目级资料块和向量索引"],
            ["实体抽取模块", "已审核笔记字段和正文", "字段优先、规则补充、术语归一化", "试剂、仪器、样本、结果等实体"],
            ["关系抽取模块", "基础实体和实验对象实体", "建立包含笔记、创建者、使用试剂、产生结果等关系", "带来源编号的一跳图谱关系"],
            ["知识图谱展示模块", "项目内实体和关系", "类型筛选、关系筛选、节点详情展示", "图谱可视化和来源跳转"],
            ["RAG 问答模块", "用户问题、资料块、可选图谱关系", "混合检索、图谱上下文注入、DeepSeek 生成", "带 S/G 证据的回答"],
            ["问答评价模块", "问答日志、审核人员评价", "准确性、可追溯性和评分记录", "可用于后续统计的评价记录"],
            ["固定任务生成模块", "已审核笔记、资料、图谱关系、任务类型", "按模板调用 DeepSeek 生成草稿", "实验总结、周报、阶段报告、图谱概览"],
            ["审计追溯模块", "业务操作、AI 调用和评价行为", "记录用户、项目、动作、目标和详情", "项目日志、AI 操作追溯链"],
        ],
    )
    insert_caption_after(p, "表 6-1 系统主要功能模块及输入输出")

    # 7.2: before/after examples at the end of the test-data subsection.
    anchor_72 = find_paragraph(doc, "7.3 功能闭环与安全边界测试", start=350)
    insert_paragraph_before(
        anchor_72,
        "为进一步说明原始实验记录如何被系统处理，表 7-2 给出三类代表性记录的“处理前—处理后”样例。处理前的内容来自系统中已审核实验笔记的字段与正文摘要；处理后不改变原始笔记，而是额外生成实体、关系、资料块和日志引用，供图谱展示、RAG 问答和审计复核使用。",
    )
    insert_caption_before(anchor_72, "表 7-2 代表性实验记录处理前后对照")
    insert_table_before(
        anchor_72,
        [
            ["实验记录样例", "处理前的记录形态", "处理后的结构化结果", "可支持的问题"],
            ["PCR 条件优化实验", "笔记标题、实验类型 PCR、样本、Taq DNA Polymerase/dNTP/MgCl2/模板 DNA、退火温度 58℃ 条带最清晰等结果描述", "生成 PCR 实验类型实体、试剂实体、样本实体和“58℃ 条带最清晰”结果实体；形成 uses_reagent、uses_sample、produces_result 关系", "使用了哪些试剂、哪些样本、58℃ 结果来自哪条笔记"],
            ["Western Blot 蛋白表达验证", "笔记正文记录处理组、目标蛋白表达变化、抗体/试剂和成像设备等信息", "抽取试剂、仪器和结果实体；将结果关系与来源笔记绑定并写入项目图谱", "使用了哪些试剂和仪器、目标蛋白表达降低来自哪条记录"],
            ["CCK-8 细胞活力检测", "资料和笔记包含检测步骤、读数要求、处理组细胞活力下降约 18% 等结论", "审核资料进入向量块；笔记产生细胞活力结果实体和相关关系；问答日志保存资料块与图谱关系", "CCK-8 步骤和读数要求来自哪份资料、18% 下降结论来自哪条笔记"],
        ],
    )

    # Appendix B: rename B.2 to B.3, insert B.2 field summary before it.
    appendix_b2 = find_paragraph(doc, "B.2 实验附件与人工盲评协议", start=500)
    set_paragraph_text(appendix_b2, "B.3 实验附件与人工盲评协议")
    appendix_b2.style = "Heading 2"
    heading = insert_paragraph_before(appendix_b2, "B.2 核心数据表字段摘要")
    heading.style = "Heading 2"
    insert_paragraph_before(
        appendix_b2,
        "表 B-2 进一步列出主要数据表的核心字段、主外键关系和字段含义。完整模型以代码中的 SQLAlchemy 定义为准，本表用于说明论文涉及的数据如何在数据库中落地。",
    )
    insert_caption_before(appendix_b2, "表 B-2 核心数据表字段摘要")
    insert_table_before(
        appendix_b2,
        [
            ["数据表", "主键", "主要外键", "核心字段", "字段含义"],
            ["users", "id", "无", "username, role, status", "保存账号、系统角色和账号状态"],
            ["projects", "id", "owner_user_id", "name, status, approval_enabled", "保存项目、负责人、状态和审批配置"],
            ["project_members", "id", "project_id, user_id", "project_role, can_read/write/review/manage", "保存项目成员角色和四类项目权限"],
            ["experiment_notes", "id", "project_id, owner_user_id", "title, experiment_type, status, current_version_id", "保存笔记归属、类型、状态和当前版本"],
            ["note_versions", "id", "note_id, created_by", "fixed_fields_json, content_json, change_summary", "保存每次笔记版本的字段、正文和变更说明"],
            ["files", "id", "project_id, note_id, uploaded_by", "file_category, storage_path, file_hash, status", "保存附件和知识文档的存储、哈希与审核状态"],
            ["rag_document_chunks", "id", "project_id, file_id", "chunk_index, content_hash, embedding, metadata_json", "保存资料分块、内容哈希、向量和元数据"],
            ["kg_entities", "id", "project_id", "entity_type, label, source_id, properties", "保存项目级实体、归一化名称、来源和属性"],
            ["kg_relations", "id", "project_id, source_entity_id, target_entity_id", "relation_type, source_id, confidence, properties", "保存实体间一跳关系、来源和置信度"],
            ["ai_query_logs", "id", "project_id, user_id, experiment_run_id", "question, answer, rag_mode, sources_json, usage_json", "保存 AI 问答、资料来源、图谱依据和 Token 信息"],
            ["agent_generation_runs", "id", "project_id, user_id", "task_type, input_params_json, body, source_ids, status", "保存固定任务生成结果、输入参数和来源编号"],
            ["audit_logs", "id", "actor_user_id, project_id", "action, target_type, detail_json, created_at", "保存关键业务操作和 AI 操作的审计信息"],
        ],
    )

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
