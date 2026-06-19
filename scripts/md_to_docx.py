"""Convert the thesis Markdown source into current Word deliverables."""

from __future__ import annotations

import json
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_LINE_SPACING,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image


ROOT = Path(__file__).resolve().parent.parent
MD_PATH = ROOT / "docs" / "毕业论文初稿.md"
DOCX_PATH = ROOT / "docs" / "毕业论文初稿.docx"
VERSIONED_PATH = ROOT / "docs" / "毕业论文初稿_v4.docx"
REVIEW_PATH = ROOT / "docs" / "毕业论文_导师审阅版.docx"
TOC_PAGE_MAP_PATH = ROOT / "docs" / "thesis_toc_pages.json"
FORMATTED_ASSET_DIR = ROOT / "tmp" / "thesis-format-assets"
COVER_MAIN_TITLE = "智能电子实验笔记系统设计与实现"
COVER_SUBTITLE = "——面向科研实验记录"
ENGLISH_TITLE = (
    "Design and Implementation of an Intelligent Electronic Laboratory "
    "Notebook System: Scientific Experiment Records"
)

FIGURES = (
    (
        "整体组织结构如图 1-1 所示",
        ROOT / "docs" / "user-guide-assets" / "01-thesis-structure.png",
        "图 1-1 论文组织结构",
    ),
    (
        "系统级与项目级角色权限边界如图 3-1 所示",
        ROOT / "docs" / "user-guide-assets" / "11-role-permission-boundary.png",
        "图 3-1 系统级与项目级角色权限边界",
    ),
    (
        "项目边界内可信知识闭环如图 3-2 所示",
        ROOT / "docs" / "user-guide-assets" / "12-trusted-knowledge-loop.png",
        "图 3-2 项目边界内可信知识闭环",
    ),
    (
        "实验知识图谱实体关系模型如图 4-1 所示",
        ROOT / "docs" / "user-guide-assets" / "13-kg-schema.png",
        "图 4-1 实验知识图谱实体关系模型",
    ),
    (
        "审核约束下的实验知识图谱构建流程如图 4-2 所示",
        ROOT / "docs" / "user-guide-assets" / "14-kg-construction-flow.png",
        "图 4-2 审核约束下的实验知识图谱构建流程",
    ),
    (
        "普通 RAG 与图谱增强 RAG 的对照流程如图 5-1 所示",
        ROOT / "docs" / "user-guide-assets" / "15-rag-comparison-flow.png",
        "图 5-1 普通 RAG 与图谱增强 RAG 对照流程",
    ),
    (
        "AI 问答证据与审计追溯链如图 5-2 所示",
        ROOT / "docs" / "user-guide-assets" / "16-traceability-chain.png",
        "图 5-2 AI 问答证据与审计追溯链",
    ),
    (
        "系统总体架构如图 6-1 所示",
        ROOT / "docs" / "user-guide-assets" / "17-system-architecture.png",
        "图 6-1 智能电子实验笔记系统总体架构",
    ),
    (
        "项目工作台截图(图 6-2)",
        ROOT / "docs" / "user-guide-assets" / "02-project-workspace.png",
        "图 6-2 项目工作台界面",
    ),
    (
        "系统用户管理界面(图 6-3)",
        ROOT / "docs" / "user-guide-assets" / "03-system-admin.png",
        "图 6-3 系统用户与全局角色管理界面",
    ),
    (
        "实验笔记管理界面(图 6-4)",
        ROOT / "docs" / "user-guide-assets" / "04-experiment-notes.png",
        "图 6-4 实验笔记管理界面",
    ),
    (
        "资料库与 AI 问答评价界面(图 6-5)",
        ROOT / "docs" / "user-guide-assets" / "05-rag-files-evaluation.png",
        "图 6-5 资料库与 AI 问答评价界面",
    ),
    (
        "图 7-1 展示知识图谱可视化界面中的",
        ROOT / "docs" / "user-guide-assets" / "06-knowledge-graph.png",
        "图 7-1 实验知识图谱可视化界面",
    ),
    (
        "图 7-2 展示答案要点证据覆盖率",
        ROOT / "docs" / "user-guide-assets" / "09-rag-objective-metrics.png",
        "图 7-2 RAG 对照实验的证据覆盖、任务完成率与响应时延",
    ),
    (
        "图 7-3 展示阈值变化下的",
        ROOT / "docs" / "user-guide-assets" / "10-kg-threshold-sensitivity.png",
        "图 7-3 图谱最低得分阈值敏感性",
    ),
    (
        "智能生成页面(图 7-4)",
        ROOT / "docs" / "user-guide-assets" / "07-agent-generation.png",
        "图 7-4 固定任务型智能辅助生成界面",
    ),
    (
        "项目日志页面(图 7-5)",
        ROOT / "docs" / "user-guide-assets" / "08-project-audit-log.png",
        "图 7-5 项目审计日志界面",
    ),
)

TABLE_CAPTIONS = (
    "表 1-1 代表性研究路线对比",
    "表 3-1 系统主要角色与权限需求",
    "表 4-1 实验知识图谱实体类型",
    "表 4-2 实验知识图谱关系类型",
    "表 7-1 实验评价指标体系",
    "表 7-2 功能闭环与安全边界测试结果",
    "表 7-3 项目知识图谱规模",
    "表 7-4 关系抽取金标准核验结果",
    "表 7-5 普通 RAG 与图谱增强 RAG 总体结果",
    "表 7-6 不同问题类型任务完成率",
    "表 7-7 图谱最低得分阈值敏感性",
    "表 7-8 固定任务型智能辅助生成验证结果",
    "表 A-1 RAG 对照实验问题清单",
    "表 B-1 核心数据表映射",
)

BASE_TOC_PAGES = {
    "摘要": 6,
    "Abstract": 7,
    "第一章 绪论": 8,
    "1.1 研究背景及意义": 8,
    "1.2 国内外研究现状": 9,
    "1.3 现有研究存在的问题": 12,
    "1.4 论文主要研究内容": 13,
    "1.5 论文创新点": 14,
    "1.6 论文组织结构": 15,
    "第二章 相关理论与技术": 15,
    "2.1 电子实验笔记与科研数据管理": 15,
    "2.2 大语言模型与检索增强生成技术": 17,
    "2.3 知识图谱与实体关系建模": 18,
    "2.4 智能体与工具调用机制": 19,
    "2.5 系统开发相关技术": 20,
    "2.6 本章小结": 21,
    "第三章 面向科研实验记录的需求分析与问题建模": 22,
    "3.1 科研实验记录业务场景分析": 22,
    "3.2 用户角色与权限需求分析": 23,
    "3.3 实验笔记、附件资料与项目知识库需求分析": 24,
    "3.4 AI 辅助实验记录管理需求分析": 25,
    "3.5 系统非功能需求分析": 25,
    "3.6 关键问题建模": 26,
    "3.7 本章小结": 28,
    "第四章 实验知识图谱构建与可信知识管理方法": 28,
    "4.1 实验知识图谱总体设计": 29,
    "4.2 实验记录实体定义": 29,
    "4.3 实验知识关系定义": 30,
    "4.4 基于实验笔记内容的实体关系抽取方法": 32,
    "4.5 面向项目权限的知识图谱更新机制": 34,
    "4.6 资料审核与知识入库机制": 34,
    "4.7 实验知识图谱可视化与关联检索": 35,
    "4.8 本章小结": 35,
    "第五章 基于知识图谱与 RAG 的智能问答及固定任务型辅助生成方法": 36,
    "5.1 项目级 RAG 知识库构建方法": 36,
    "5.2 融合知识图谱的检索增强策略": 37,
    "5.3 权限约束下的智能问答流程": 39,
    "5.4 引用来源与回答可追溯机制": 39,
    "5.5 面向实验笔记的固定任务型辅助生成设计": 41,
    "5.6 基于工具调用/MCP 的辅助生成机制": 42,
    "5.7 本章小结": 43,
    "第六章 智能电子实验笔记系统设计与实现": 43,
    "6.1 系统架构设计": 43,
    "6.2 系统功能模块设计": 44,
    "6.3 数据库设计": 46,
    "6.4 用户与权限模块实现": 46,
    "6.5 项目与实验笔记模块实现": 47,
    "6.6 附件资料与审核模块实现": 48,
    "6.7 知识图谱模块实现": 49,
    "6.8 AI 问答与智能辅助生成模块实现": 49,
    "6.9 安全与审计追溯实现": 50,
    "6.10 本章小结": 51,
    "第七章 实验与结果分析": 51,
    "7.1 实验目标与评价指标": 51,
    "7.2 实验环境与测试数据": 53,
    "7.3 功能闭环与安全边界测试": 54,
    "7.4 知识图谱构建效果分析": 55,
    "7.5 普通 RAG 与图谱增强 RAG 对照实验": 57,
    "7.6 固定任务型智能辅助生成验证": 60,
    "7.7 场景化验证与审计追溯": 61,
    "7.8 实验结果分析": 62,
    "7.9 本章小结": 63,
    "第八章 总结与展望": 64,
    "8.1 工作总结": 64,
    "8.2 主要创新点总结": 64,
    "8.3 不足分析": 65,
    "8.4 未来展望": 65,
    "附录 A RAG 对照实验问题清单": 65,
    "附录 B 系统实现与评价补充材料": 67,
    "B.1 核心数据表映射": 67,
    "B.2 实验附件与人工盲评协议": 68,
    "参考文献": 69,
}
TOC_ENTRIES_PER_PAGE = 30
TOC_BODY_PAGE_SHIFT = 0


def set_run_font(run, chinese: str = "宋体", western: str = "Times New Roman", size: int = 12) -> None:
    run.font.name = western
    run._element.rPr.rFonts.set(qn("w:eastAsia"), chinese)
    run.font.size = Pt(size)


def set_exact_line_spacing(paragraph_format, points: int) -> None:
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph_format.line_spacing = Pt(points)


def set_first_line_chars(paragraph_format, chars: int) -> None:
    paragraph_properties = paragraph_format._element.get_or_add_pPr()
    indent = paragraph_properties.get_or_add_ind()
    for attribute in ("firstLine", "hanging", "hangingChars"):
        indent.attrib.pop(qn(f"w:{attribute}"), None)
    indent.set(qn("w:firstLineChars"), str(chars * 100))


def configure_section(section) -> None:
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.8)
    section.bottom_margin = Cm(2.8)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)


def add_page_number(section, start: int = 1) -> None:
    section.footer.is_linked_to_previous = False
    footer = section.footer.paragraphs[0]
    footer.clear()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.first_line_indent = Cm(0)
    run = footer.add_run()
    set_run_font(run, size=10)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = str(start)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])

    page_number_type = section._sectPr.find(qn("w:pgNumType"))
    if page_number_type is None:
        page_number_type = OxmlElement("w:pgNumType")
        section._sectPr.append(page_number_type)
    page_number_type.set(qn("w:start"), str(start))
    page_number_type.set(qn("w:fmt"), "decimal")


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    configure_section(section)
    section.footer.is_linked_to_previous = False
    section.footer.paragraphs[0].clear()

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    set_exact_line_spacing(normal.paragraph_format, 20)
    set_first_line_chars(normal.paragraph_format, 2)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    heading_specs = {
        "Title": ("宋体", 22, WD_ALIGN_PARAGRAPH.CENTER),
        "Heading 1": ("宋体", 18, WD_ALIGN_PARAGRAPH.CENTER),
        "Heading 2": ("黑体", 14, WD_ALIGN_PARAGRAPH.LEFT),
        "Heading 3": ("黑体", 12, WD_ALIGN_PARAGRAPH.LEFT),
    }
    for style_name, (font_name, size, alignment) in heading_specs.items():
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.alignment = alignment
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.first_line_indent = Cm(0)
        set_exact_line_spacing(style.paragraph_format, 20)

    zoom = doc.settings.element.find(qn("w:zoom"))
    if zoom is not None:
        zoom.set(qn("w:percent"), "100")


def add_centered_text(
    doc: Document,
    text: str,
    *,
    chinese: str = "宋体",
    size: int = 12,
    bold: bool = False,
    before: int = 0,
    after: int = 0,
) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    run = paragraph.add_run(text)
    run.bold = bold
    set_run_font(run, chinese=chinese, size=size)


def add_cover(doc: Document, title: str) -> None:
    classification = doc.add_paragraph()
    classification.alignment = WD_ALIGN_PARAGRAPH.LEFT
    classification.paragraph_format.first_line_indent = Cm(0)
    run = classification.add_run("分类号：____________")
    set_run_font(run, size=12)

    add_centered_text(doc, "安徽师范大学", chinese="宋体", size=22, bold=True, before=36)
    add_centered_text(doc, "硕士学位论文", chinese="宋体", size=22, bold=True, after=28)
    add_centered_text(
        doc,
        COVER_MAIN_TITLE,
        chinese="宋体",
        size=22,
        bold=True,
        before=24,
        after=4,
    )
    add_centered_text(
        doc,
        COVER_SUBTITLE,
        chinese="宋体",
        size=16,
        bold=True,
        after=12,
    )
    add_centered_text(doc, ENGLISH_TITLE, size=14, bold=True, after=36)

    for label in ("学科专业", "研究方向", "作者姓名", "指导教师", "论文提交日期"):
        add_centered_text(doc, f"{label}：____________________", size=14, after=8)


def add_declarations(doc: Document) -> None:
    add_centered_text(doc, "学位论文独创性声明", chinese="宋体", size=18, bold=True, after=18)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_exact_line_spacing(paragraph.paragraph_format, 20)
    set_first_line_chars(paragraph.paragraph_format, 2)
    add_inline_runs(
        paragraph,
        "本人郑重声明：所呈交的学位论文是在指导教师指导下独立完成的研究成果。"
        "除文中已经注明引用的内容外，本论文不包含他人已经发表或撰写过的研究成果。"
        "对本文研究作出重要贡献的个人和集体，均已在文中明确说明并表示谢意。",
    )
    add_centered_text(doc, "作者签名：____________    日期：____年__月__日", before=28)

    add_centered_text(
        doc,
        "学位论文版权使用授权书",
        chinese="宋体",
        size=18,
        bold=True,
        before=48,
        after=18,
    )
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_exact_line_spacing(paragraph.paragraph_format, 20)
    set_first_line_chars(paragraph.paragraph_format, 2)
    add_inline_runs(
        paragraph,
        "本人及指导教师同意学校保留并向有关机构送交本学位论文的复印件和电子版，"
        "允许论文被查阅和借阅，并授权学校将本论文的全部或部分内容编入有关数据库"
        "进行检索、保存和汇编。涉密论文在解密后适用本授权书。",
    )
    add_centered_text(
        doc,
        "作者签名：____________    指导教师签名：____________",
        before=28,
    )
    add_centered_text(doc, "日期：____年__月__日")


def reorder_back_matter(md_text: str) -> str:
    blocks = re.split(r"(?=^##\s+)", md_text, flags=re.MULTILINE)
    prefix, content_blocks = blocks[0], blocks[1:]
    appendices = [block for block in content_blocks if block.startswith("## 附录")]
    references = [block for block in content_blocks if block.startswith("## 参考文献")]
    acknowledgements = [block for block in content_blocks if block.startswith("## 致谢")]
    main_blocks = [
        block
        for block in content_blocks
        if block not in appendices + references + acknowledgements
    ]
    return prefix + "".join(main_blocks + references + acknowledgements + appendices)


def add_toc(doc: Document, md_text: str) -> None:
    page_map = BASE_TOC_PAGES
    if TOC_PAGE_MAP_PATH.exists():
        page_map = json.loads(TOC_PAGE_MAP_PATH.read_text(encoding="utf-8"))

    headings = []
    for line in md_text.splitlines():
        match = re.match(r"^(#{2,4})\s+(.+)$", line.strip())
        if match and match.group(2) not in {"摘要", "Abstract"}:
            headings.append((len(match.group(1)), match.group(2)))

    doc.add_page_break()
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run("目录")
    run.bold = True
    set_run_font(run, chinese="宋体", size=18)

    for index, (level, title) in enumerate(headings):
        if index and index % TOC_ENTRIES_PER_PAGE == 0:
            doc.add_page_break()
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.left_indent = Cm(max(0, level - 2) * 0.74)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        set_exact_line_spacing(paragraph.paragraph_format, 16)
        paragraph.paragraph_format.tab_stops.add_tab_stop(
            Cm(15.0),
            WD_TAB_ALIGNMENT.RIGHT,
            WD_TAB_LEADER.DOTS,
        )
        page = page_map.get(title, BASE_TOC_PAGES.get(title, 1)) + TOC_BODY_PAGE_SHIFT
        run = paragraph.add_run(f"{title}\t{page}")
        run.bold = level == 2
        set_run_font(
            run,
            chinese="黑体" if level == 2 else "宋体",
            size=10 if level == 2 else 9,
        )


def start_numbered_body_section(doc: Document) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section)
    add_page_number(section, start=1)


def add_inline_runs(paragraph, text: str) -> None:
    text = text.replace("`", "")
    for part in re.split(r"(\*\*.*?\*\*)", text):
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        content = part[2:-2] if bold else part
        run = paragraph.add_run(content)
        run.bold = bold
        set_run_font(run)


def set_cell_border(cell, **edges) -> None:
    cell_properties = cell._tc.get_or_add_tcPr()
    borders = cell_properties.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        cell_properties.append(borders)
    for edge_name, edge_settings in edges.items():
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        for key, value in edge_settings.items():
            edge.set(qn(f"w:{key}"), str(value))


def add_table(doc: Document, lines: list[str], caption: str) -> None:
    rows = []
    for line in lines:
        if re.fullmatch(r"\|[\s:|-]+\|", line.strip()):
            continue
        rows.append([cell.strip().replace("`", "") for cell in line.strip().strip("|").split("|")])
    if not rows:
        return

    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.first_line_indent = Cm(0)
    caption_paragraph.paragraph_format.keep_with_next = True
    set_exact_line_spacing(caption_paragraph.paragraph_format, 15)
    caption_run = caption_paragraph.add_run(caption)
    caption_run.bold = True
    set_run_font(caption_run, size=10)

    width = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=width)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_index, row in enumerate(rows):
        for column_index in range(width):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            border_settings = {
                "top": {"val": "nil"},
                "bottom": {"val": "nil"},
                "left": {"val": "nil"},
                "right": {"val": "nil"},
                "insideH": {"val": "nil"},
                "insideV": {"val": "nil"},
            }
            if row_index == 0:
                border_settings["top"] = {"val": "single", "sz": "12", "color": "000000"}
                border_settings["bottom"] = {"val": "single", "sz": "6", "color": "000000"}
            if row_index == len(rows) - 1:
                border_settings["bottom"] = {"val": "single", "sz": "12", "color": "000000"}
            set_cell_border(cell, **border_settings)
            cell.text = row[column_index] if column_index < len(row) else ""
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.first_line_indent = Cm(0)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                set_exact_line_spacing(paragraph.paragraph_format, 15)
                for run in paragraph.runs:
                    run.bold = row_index == 0
                    set_run_font(run, chinese="黑体" if row_index == 0 else "宋体", size=9)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)


def prepare_figure_image(image_path: Path) -> Path:
    FORMATTED_ASSET_DIR.mkdir(parents=True, exist_ok=True)
    output_path = FORMATTED_ASSET_DIR / image_path.name
    with Image.open(image_path) as image:
        required_width = 1800
        if image.width < required_width:
            scale = required_width / image.width
            new_size = (required_width, round(image.height * scale))
            image = image.resize(new_size, Image.Resampling.LANCZOS)
        image.save(output_path, dpi=(300, 300))
    return output_path


def add_figure(doc: Document, image_path: Path, caption: str) -> None:
    if not image_path.exists():
        raise FileNotFoundError(image_path)
    prepared_image = prepare_figure_image(image_path)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.page_break_before = True
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.add_run().add_picture(str(prepared_image), width=Cm(14.5))

    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.first_line_indent = Cm(0)
    caption_paragraph.paragraph_format.keep_together = True
    set_exact_line_spacing(caption_paragraph.paragraph_format, 15)
    run = caption_paragraph.add_run(caption)
    run.bold = True
    set_run_font(run, size=10)


def add_algorithm(doc: Document, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""

    title = cell.paragraphs[0]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = Cm(0)
    title_run = title.add_run(lines[0])
    title_run.bold = True
    set_run_font(title_run, chinese="黑体", size=10)

    for line in lines[1:]:
        paragraph = cell.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.left_indent = Cm(0.2)
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run(line)
        set_run_font(run, western="Courier New", size=9)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)


def add_math_run(parent, text: str, plain: bool = False, align: bool = False) -> None:
    run = OxmlElement("m:r")
    if plain or align:
        run_properties = OxmlElement("m:rPr")
        if plain:
            style = OxmlElement("m:sty")
            style.set(qn("m:val"), "p")
            run_properties.append(style)
        if align:
            run_properties.append(OxmlElement("m:aln"))
        run.append(run_properties)
    value = OxmlElement("m:t")
    if text[:1].isspace() or text[-1:].isspace():
        value.set(qn("xml:space"), "preserve")
    value.text = text
    run.append(value)
    parent.append(run)


def add_math_subscript(parent, base: str, subscript: str, *, plain_base: bool = False) -> None:
    sub = OxmlElement("m:sSub")
    sub.append(OxmlElement("m:sSubPr"))
    expression = OxmlElement("m:e")
    add_math_run(expression, base, plain=plain_base)
    subscript_element = OxmlElement("m:sub")
    add_math_run(subscript_element, subscript)
    sub.extend([expression, subscript_element])
    parent.append(sub)


def add_math_fraction(parent, numerator_builder, denominator_builder) -> None:
    fraction = OxmlElement("m:f")
    fraction.append(OxmlElement("m:fPr"))
    numerator = OxmlElement("m:num")
    numerator_builder(numerator)
    denominator = OxmlElement("m:den")
    denominator_builder(denominator)
    fraction.extend([numerator, denominator])
    parent.append(fraction)


def add_math_sum(parent, lower_builder, upper_text: str | None, expression_builder) -> None:
    summation = OxmlElement("m:nary")
    properties = OxmlElement("m:naryPr")
    character = OxmlElement("m:chr")
    character.set(qn("m:val"), "∑")
    limit_location = OxmlElement("m:limLoc")
    limit_location.set(qn("m:val"), "undOvr")
    properties.extend([character, limit_location])
    if upper_text is None:
        hide_upper = OxmlElement("m:supHide")
        hide_upper.set(qn("m:val"), "1")
        properties.append(hide_upper)
    lower = OxmlElement("m:sub")
    lower_builder(lower)
    upper = OxmlElement("m:sup")
    if upper_text is not None:
        add_math_run(upper, upper_text)
    expression = OxmlElement("m:e")
    expression_builder(expression)
    summation.extend([properties, lower, upper, expression])
    parent.append(summation)


def add_equation_table(doc: Document, equation_number: str):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = (Cm(0.5), Cm(12.0), Cm(2.0))
    for column, width in zip(table.columns, widths):
        column.width = width
    for cell, width in zip(table.rows[0].cells, widths):
        cell.width = width
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    equation_paragraph = table.cell(0, 1).paragraphs[0]
    equation_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    equation_paragraph.paragraph_format.first_line_indent = Cm(0)
    equation_paragraph.paragraph_format.line_spacing = 1.0

    number_paragraph = table.cell(0, 2).paragraphs[0]
    number_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    number_paragraph.paragraph_format.first_line_indent = Cm(0)
    number_cell_properties = table.cell(0, 2)._tc.get_or_add_tcPr()
    number_cell_properties.append(OxmlElement("w:noWrap"))
    cell_margins = OxmlElement("w:tcMar")
    for side in ("top", "left", "bottom", "right"):
        margin = OxmlElement(f"w:{side}")
        margin.set(qn("w:w"), "0")
        margin.set(qn("w:type"), "dxa")
        cell_margins.append(margin)
    number_cell_properties.append(cell_margins)
    number = number_paragraph.add_run(f"({equation_number})")
    set_run_font(number, size=10)
    return equation_paragraph


def add_known_equation(doc: Document, equation_number: str) -> None:
    paragraph = add_equation_table(doc, equation_number)
    math = OxmlElement("m:oMath")

    if equation_number == "4-1":
        delimiter = OxmlElement("m:d")
        delimiter_properties = OxmlElement("m:dPr")
        begin_character = OxmlElement("m:begChr")
        begin_character.set(qn("m:val"), "{")
        end_character = OxmlElement("m:endChr")
        end_character.set(qn("m:val"), "")
        delimiter_properties.extend([begin_character, end_character])
        delimiter.append(delimiter_properties)
        delimiter_expression = OxmlElement("m:e")
        matrix = OxmlElement("m:m")
        matrix_properties = OxmlElement("m:mPr")
        matrix_columns = OxmlElement("m:mcs")
        for alignment_value in ("right", "left"):
            matrix_column = OxmlElement("m:mc")
            matrix_column_properties = OxmlElement("m:mcPr")
            column_count = OxmlElement("m:count")
            column_count.set(qn("m:val"), "1")
            column_alignment = OxmlElement("m:mcJc")
            column_alignment.set(qn("m:val"), alignment_value)
            matrix_column_properties.extend([column_count, column_alignment])
            matrix_column.append(matrix_column_properties)
            matrix_columns.append(matrix_column)
        matrix_properties.append(matrix_columns)
        matrix.append(matrix_properties)
        for value, condition in (
            ("1.0", "r 来自数据库元数据"),
            ("0.7", "r 来自结构化字段或正文规则抽取"),
        ):
            row = OxmlElement("m:mr")
            value_cell = OxmlElement("m:e")
            add_math_run(value_cell, value)
            condition_cell = OxmlElement("m:e")
            add_math_run(condition_cell, f",  {condition}", plain=True)
            row.extend([value_cell, condition_cell])
            matrix.append(row)
        delimiter_expression.append(matrix)
        delimiter.append(delimiter_expression)
        add_math_run(math, "conf", plain=True)
        add_math_run(math, "(r) = ")
        math.append(delimiter)
    elif equation_number == "5-1":
        add_math_subscript(math, "R", "i")
        add_math_run(math, " = 0.8")
        add_math_subscript(math, "v", "i")
        add_math_run(math, " + 0.2")
        add_math_subscript(math, "l", "i")
    elif equation_number == "5-2":
        matrix = OxmlElement("m:m")
        matrix_properties = OxmlElement("m:mPr")
        matrix_columns = OxmlElement("m:mcs")
        matrix_column = OxmlElement("m:mc")
        matrix_column_properties = OxmlElement("m:mcPr")
        column_count = OxmlElement("m:count")
        column_count.set(qn("m:val"), "1")
        column_alignment = OxmlElement("m:mcJc")
        column_alignment.set(qn("m:val"), "left")
        matrix_column_properties.extend([column_count, column_alignment])
        matrix_column.append(matrix_column_properties)
        matrix_columns.append(matrix_column)
        matrix_properties.append(matrix_columns)
        matrix.append(matrix_properties)

        first_row = OxmlElement("m:mr")
        first_expression = OxmlElement("m:e")
        add_math_run(first_expression, "G")
        add_math_run(first_expression, "(r,q) = 3")
        add_math_subscript(first_expression, "I", "type")
        add_math_run(first_expression, " + ")

        def lower_sum(parent):
            add_math_run(parent, "k∈K(q)")

        def graph_term(parent):
            add_math_run(parent, "(3")
            add_math_subscript(parent, "I", "exact")
            add_math_run(parent, " + ")
            add_math_subscript(parent, "I", "substring")
            add_math_run(parent, ")")

        add_math_sum(first_expression, lower_sum, None, graph_term)
        first_row.append(first_expression)
        matrix.append(first_row)

        second_row = OxmlElement("m:mr")
        second_expression = OxmlElement("m:e")
        add_math_run(second_expression, "                         + 0.2")
        add_math_subscript(second_expression, "I", "note")
        add_math_run(second_expression, " + 0.3")
        add_math_subscript(second_expression, "I", "extraction")
        second_row.append(second_expression)
        matrix.append(second_row)
        math.append(matrix)
    elif equation_number == "7-1":
        add_math_run(math, "P = ")
        add_math_fraction(
            math,
            lambda parent: add_math_run(parent, "TP"),
            lambda parent: (
                add_math_run(parent, "TP + FP")
            ),
        )
        add_math_run(math, ",    R = ")
        add_math_fraction(
            math,
            lambda parent: add_math_run(parent, "TP"),
            lambda parent: add_math_run(parent, "TP + FN"),
        )
    elif equation_number == "7-2":
        add_math_subscript(math, "F", "1")
        add_math_run(math, " = ")
        add_math_fraction(
            math,
            lambda parent: add_math_run(parent, "2PR"),
            lambda parent: add_math_run(parent, "P + R"),
        )
    elif equation_number in {"7-3", "7-4", "7-5"}:
        left_labels = {
            "7-3": ("C", None, "c"),
            "7-4": ("Src", "avg", "src"),
            "7-5": ("T", "avg", "ms"),
        }
        base, subscript, term = left_labels[equation_number]
        if subscript:
            add_math_subscript(math, base, subscript, plain_base=True)
        else:
            add_math_run(math, base)
        add_math_run(math, " = ")
        add_math_fraction(
            math,
            lambda parent: add_math_run(parent, "1"),
            lambda parent: add_math_run(parent, "N"),
        )
        add_math_run(math, " × ")

        def lower_sum(parent):
            add_math_run(parent, "i = 1")

        def average_term(parent):
            add_math_subscript(parent, term, "i", plain_base=term in {"src", "hit", "ms"})

        add_math_sum(math, lower_sum, "N", average_term)
    else:
        raise ValueError(f"Unsupported equation number: {equation_number}")

    paragraph._p.append(math)


def add_permission_equation(doc: Document) -> None:
    equation_paragraph = add_equation_table(doc, "6-1")

    math = OxmlElement("m:oMath")
    delimiter = OxmlElement("m:d")
    delimiter_properties = OxmlElement("m:dPr")
    begin_character = OxmlElement("m:begChr")
    begin_character.set(qn("m:val"), "{")
    end_character = OxmlElement("m:endChr")
    end_character.set(qn("m:val"), "")
    delimiter_properties.extend([begin_character, end_character])
    delimiter.append(delimiter_properties)

    delimiter_expression = OxmlElement("m:e")
    matrix = OxmlElement("m:m")
    matrix_properties = OxmlElement("m:mPr")
    matrix_columns = OxmlElement("m:mcs")
    for alignment_value in ("right", "center", "left"):
        matrix_column = OxmlElement("m:mc")
        matrix_column_properties = OxmlElement("m:mcPr")
        column_count = OxmlElement("m:count")
        column_count.set(qn("m:val"), "1")
        column_alignment = OxmlElement("m:mcJc")
        column_alignment.set(qn("m:val"), alignment_value)
        matrix_column_properties.extend([column_count, column_alignment])
        matrix_column.append(matrix_column_properties)
        matrix_columns.append(matrix_column)
    matrix_properties.append(matrix_columns)
    matrix.append(matrix_properties)

    equations = (
        ("Read", ("S(u)", "O(u, p)", "M(u, p)")),
        ("Write", ("S(u)", "W(u, p)")),
        ("Review", ("S(u)", "V(u, p)", "G(u, p)")),
        ("Manage", ("S(u)", "O(u, p)", "G(u, p)")),
    )
    for function_name, terms in equations:
        matrix_row = OxmlElement("m:mr")
        function_cell = OxmlElement("m:e")
        add_math_run(function_cell, function_name, plain=True)
        add_math_run(function_cell, "(u, p)")
        matrix_row.append(function_cell)

        equals_cell = OxmlElement("m:e")
        add_math_run(equals_cell, "=")
        matrix_row.append(equals_cell)

        expression_cell = OxmlElement("m:e")
        for term_index, term in enumerate(terms):
            if term_index:
                add_math_run(expression_cell, " OR ", plain=True)
            add_math_run(expression_cell, term)
        matrix_row.append(expression_cell)
        matrix.append(matrix_row)

    delimiter_expression.append(matrix)
    delimiter.append(delimiter_expression)
    math.append(delimiter)
    equation_paragraph._p.append(math)

def build_document(md_text: str) -> Document:
    md_text = reorder_back_matter(md_text)
    doc = Document()
    configure_document(doc)

    lines = md_text.splitlines()
    index = 0
    title_written = False
    body_started = False
    english_abstract = False
    table_index = 0
    inserted_figures: set[Path] = set()
    while index < len(lines):
        raw = lines[index]
        line = raw.strip()
        if not line:
            index += 1
            continue

        if not title_written:
            add_cover(doc, line)
            doc.add_page_break()
            add_declarations(doc)
            doc.add_page_break()
            title_written = True
            index += 1
            continue

        if line.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            if table_index >= len(TABLE_CAPTIONS):
                raise ValueError("More Markdown tables than configured table captions")
            add_table(doc, table_lines, TABLE_CAPTIONS[table_index])
            table_index += 1
            continue

        if re.match(r"^算法\s+\d+-\d+\s+", line):
            algorithm_lines = [line]
            index += 1
            while index < len(lines) and lines[index].strip():
                candidate = lines[index].strip()
                if candidate.startswith("#") or candidate.startswith("|"):
                    break
                algorithm_lines.append(candidate)
                index += 1
            add_algorithm(doc, algorithm_lines)
            continue

        if line.startswith("Read(u, p) ="):
            equation_lines = lines[index : index + 4]
            if len(equation_lines) != 4 or not equation_lines[-1].strip().endswith("(6-1)"):
                raise ValueError("Malformed permission equation block")
            add_permission_equation(doc)
            index += 4
            continue

        if line.startswith("conf(r) = 1.0"):
            equation_lines = lines[index : index + 2]
            if len(equation_lines) != 2 or not equation_lines[-1].strip().endswith("(式4-1)"):
                raise ValueError("Malformed confidence equation block")
            add_known_equation(doc, "4-1")
            index += 2
            continue

        heading_match = re.match(r"^(#{2,4})\s+(.+)$", line)
        if heading_match:
            heading_title = heading_match.group(2)
            if heading_title == "Abstract":
                doc.add_page_break()
                english_abstract = True
            elif heading_title == "第一章 绪论":
                english_abstract = False
                add_toc(doc, md_text)
                start_numbered_body_section(doc)
                body_started = True
            elif len(heading_match.group(1)) == 2 and heading_title != "摘要":
                doc.add_page_break()
            level = len(heading_match.group(1)) - 1
            heading = doc.add_heading(heading_title, level=level)
            heading.paragraph_format.first_line_indent = Cm(0)
            index += 1
            continue

        if re.match(r"^\d+\.\s+", line):
            paragraph = doc.add_paragraph(style="List Number")
            paragraph.paragraph_format.first_line_indent = Cm(0)
            add_inline_runs(paragraph, re.sub(r"^\d+\.\s+", "", line))
        elif line.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.first_line_indent = Cm(0)
            add_inline_runs(paragraph, line[2:])
        else:
            equation_match = re.match(r"^(.*?)(?:\s+\(式(\d+-\d+)\))$", line)
            if equation_match:
                add_known_equation(doc, equation_match.group(2))
            else:
                paragraph = doc.add_paragraph()
                add_inline_runs(paragraph, line)
                if english_abstract:
                    set_exact_line_spacing(paragraph.paragraph_format, 18)

        for trigger, image_path, caption in FIGURES:
            if trigger in line and image_path not in inserted_figures:
                add_figure(doc, image_path, caption)
                inserted_figures.add(image_path)
                break
        index += 1

    if not body_started:
        raise ValueError("Unable to find the first chapter and start body page numbering")
    if table_index != len(TABLE_CAPTIONS):
        raise ValueError(
            f"Configured {len(TABLE_CAPTIONS)} table captions but rendered {table_index} tables"
        )
    return doc


def main() -> None:
    doc = build_document(MD_PATH.read_text(encoding="utf-8"))
    doc.save(DOCX_PATH)
    shutil.copy2(DOCX_PATH, VERSIONED_PATH)
    shutil.copy2(DOCX_PATH, REVIEW_PATH)
    for path in (DOCX_PATH, VERSIONED_PATH, REVIEW_PATH):
        print(f"Word document saved: {path} ({path.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
